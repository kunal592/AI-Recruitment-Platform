"""
app/resume_parser/extractor.py
─────────────────────────────────────────────────────────────────────────────
Extracts raw text from uploaded PDF or DOCX resume files.
Uses pdfplumber for PDFs and python-docx for Word documents.
"""

import io
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from loguru import logger


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from a PDF file.

    Args:
        file_bytes: Raw PDF bytes.

    Returns:
        Concatenated text from all pages.
    """
    text_parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    logger.debug("PDF page {} extracted {} chars", page_num, len(page_text))
    except Exception as exc:
        logger.error("PDF extraction failed: {}", exc)
        raise ValueError(f"Could not extract text from PDF: {exc}") from exc

    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file.

    Args:
        file_bytes: Raw DOCX bytes.

    Returns:
        Concatenated paragraph text.
    """
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        logger.error("DOCX extraction failed: {}", exc)
        raise ValueError(f"Could not extract text from DOCX: {exc}") from exc


def extract_resume_text(file_bytes: bytes, file_extension: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.

    Args:
        file_bytes:     Raw file bytes.
        file_extension: Lowercase extension without dot — "pdf" or "docx".

    Returns:
        Extracted plain text.

    Raises:
        ValueError: If the file type is unsupported or extraction fails.
    """
    ext = file_extension.lower().lstrip(".")
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Only PDF and DOCX are accepted.")
